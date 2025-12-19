"""Word report generator for MIT/VRI analytics."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import date
from typing import Iterable, Optional

import matplotlib
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from .clickhouse_io import CH
from .utils import normalize_manufacturer_name

matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402


FIXED_PRICES = {
    "57696-14": 14000,
    "54566-13": 14000,
    "53219-13": 27000,
    "73260-18": 27000,
    "69124-17": 27000,
    "71897-18": 20700,
    "65740-16": 18000,
    "50927-12": 20700,
    "51122-12": 14100,
    "58588-14": 20700,
    "51526-12": 12500,
    "90029-23": 12500,
    "90006-23": 12500,
    "64052-16": 12500,
    "86719-22": 12500,
    "27903-04": 8000,
    "51117-12": 8000,
    "51116-12": 10000,
    "81429-21": 10000,
    "82561-21": 10000,
    "64303-16": 10000,
    "86424-22": 10000,
    "56869-14": 10000,
    "63734-16": 10000,
    "84581-22": 10000,
    "84676-22": 10000,
    "64333-16": 12500,
    "86348-22": 12500,
    "27940-04": 12500,
    "70594-18": 12500,
    "68425-17": 12500,
    "84413-22": 12500,
    "51521-12": 8000,
    "78227-20": 8000,
    "51523-12": 12500,
    "86357-22": 12500,
    "59837-15": 8000,
    "76291-19": 8000,
    "66993-17": 8000,
    "95403-25": 8000,
    "51524-12": 13000,
    "64268-16": 11000,
    "82090-21": 11000,
    "86126-22": 11000,
    "96698-25": 11000,
    "96997-25": 11000,
    "96998-25": 10000,
    "86815-22": 56000,
    "88826-23": 21000,
    "78058-20": 37000,
    "82894-21": 37000,
    "87478-22": 10000,
    "31760-06": 139460,
    "59304-14": 124680,
    "59303-14": 1154780,
    "64053-16": 221220,
    "67305-17": 173845,
    "91158-24": 65800,
    "70075-17": 103300,
    "71891-18": 79000,
    "93163-24": 58900,
    "76758-19": 93640,
    "70805-18": 141618,
    "77286-20": 141618,
    "60693-15": 49200,
    "78552-20": 55200,
    "86888-22": 355200,
    "57573-14": 261650,
    "54122-13": 199200,
    "86891-22": 215600,
    "96878-25": 215600,
    "85996-22": 56600,
    "51643-12": 168400,
    "88354-23": 600000,
    "83775-21": 143800,
    "85912-22": 143800,
    "94918-25": 20700,
    "88767-23": 14000,
    "95442-25": 27000,
    "83082-21": 76800,
    "91055-24": 50000,
    "86245-22": 630000,
    "56957-14": 38400,
}
COMPLEX_SYSTEMS = ["45982-10", "64868-16", "75380-19", "96098-25"]


@dataclass
class ReportParams:
    year_from: int
    year_to: int
    our_group_name: str
    our_org_regex: str
    top_n: int = 20
    competitors_n: int = 10
    only_applicable: bool = True
    serial_or_us: bool = True
    mit_numbers: Optional[Iterable[str]] = None


def _mit_year_expr() -> str:
    part = "splitByChar('-', mit_number)[-1]"
    year_num = f"toInt32OrZero({part})"
    return (
        f"if(isNotNull(order_date), toYear(order_date), "
        f"if(length({part}) = 2, 2000 + {year_num}, if(length({part}) = 4, {year_num}, 0)))"
    )


def _set_cell_shading(cell, fill: str = "000000") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_df_table(
    doc: Document,
    df: pd.DataFrame,
    title: str,
    *,
    highlight_col: Optional[str] = None,
    highlight_value: Optional[str] = None,
    int_cols: Optional[list[str]] = None,
    money_cols: Optional[list[str]] = None,
) -> None:
    doc.add_heading(title, level=2)
    df2 = df.copy()
    int_cols = int_cols or []
    money_cols = money_cols or []

    for col in int_cols:
        if col in df2.columns:
            df2[col] = df2[col].apply(lambda x: "" if pd.isna(x) else f"{int(x):,}".replace(",", " "))
    for col in money_cols:
        if col in df2.columns:
            df2[col] = df2[col].apply(lambda x: "" if pd.isna(x) else f"{int(x):,}".replace(",", " "))

    include_index = df2.index.name is not None
    cols = ([df2.index.name] if include_index else []) + list(df2.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"

    for j, col in enumerate(cols):
        table.rows[0].cells[j].text = str(col)

    for i in range(len(df2)):
        row_cells = table.add_row().cells
        highlight = False
        if highlight_col and highlight_col in df2.columns:
            highlight = str(df2.iloc[i][highlight_col]) == str(highlight_value)

        k = 0
        if include_index:
            row_cells[k].text = str(df2.index[i])
            k += 1
        for col in df2.columns:
            row_cells[k].text = str(df2.iloc[i][col])
            k += 1

        if highlight:
            for cell in row_cells:
                _set_cell_shading(cell, fill="000000")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = None

    doc.add_paragraph("")


def _calc_verification_price(mit_number: str, mi_modification: str) -> int:
    if mit_number in FIXED_PRICES:
        return FIXED_PRICES[mit_number]
    if mit_number in COMPLEX_SYSTEMS:
        if not mi_modification or not isinstance(mi_modification, str):
            return 37000
        parts = re.findall(r"\d+", mi_modification)
        ch = 0
        if mit_number in ["64868-16", "75380-19", "96098-25"] and len(parts) >= 3:
            ch = int(parts[-2])
        elif mit_number == "45982-10" and len(parts) >= 1:
            ch = int(parts[-1])
        if ch > 0:
            per_ch = 33 if ch <= 1000 else 22 if ch <= 4000 else 14 if ch <= 10000 else 12
            return 37000 + (ch * per_ch)
        return 37000
    return 0


def _fig_to_bytes(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _fetch_mit_counts(ch: CH, year_from: int, year_to: int) -> pd.DataFrame:
    year_expr = _mit_year_expr()
    sql = (
        "SELECT "
        f"{year_expr} AS year, manufacturer, production_type, count() AS cnt "
        f"FROM {ch.db}.mit_registry "
        f"WHERE {year_expr} BETWEEN {int(year_from)} AND {int(year_to)} AND {year_expr} > 0 "
        "GROUP BY year, manufacturer, production_type"
    )
    rows = ch.rows(sql)
    return pd.DataFrame(rows, columns=["year", "manufacturer", "production_type", "count"])


def _fetch_vri_agg(
    ch: CH,
    year_from: int,
    year_to: int,
    only_applicable: bool,
    mit_numbers: Iterable[str],
) -> pd.DataFrame:
    numbers = [str(x) for x in mit_numbers if x]
    if not numbers:
        return pd.DataFrame(columns=["year", "mit_number", "mi_modification", "org_title", "qty"])
    numbers_sql = "', '".join(numbers)
    where = [f"toYear(verification_date) BETWEEN {int(year_from)} AND {int(year_to)}"]
    if only_applicable:
        where.append("applicability = 1")
    where.append(f"mit_number IN ('{numbers_sql}')")
    sql = (
        "SELECT "
        "toYear(verification_date) AS year, "
        "mit_number, mi_modification, org_title, "
        "countDistinct(vri_id) AS qty "
        f"FROM {ch.db}.verifications "
        f"WHERE {' AND '.join(where)} "
        "GROUP BY year, mit_number, mi_modification, org_title"
    )
    rows = ch.rows(sql)
    return pd.DataFrame(rows, columns=["year", "mit_number", "mi_modification", "org_title", "qty"])


def build_word_report(ch: CH, params: ReportParams) -> bytes:
    year_from = min(params.year_from, params.year_to)
    year_to = max(params.year_from, params.year_to)
    mit_numbers = params.mit_numbers or list(FIXED_PRICES.keys()) + list(COMPLEX_SYSTEMS)

    df_mit = _fetch_mit_counts(ch, year_from, year_to)
    df_mit["short_name"] = df_mit["manufacturer"].apply(normalize_manufacturer_name)
    df_mit["year"] = df_mit["year"].fillna(0).astype(int)

    if params.serial_or_us:
        mask_serial_or_us = (df_mit["production_type"] == 1) | (df_mit["short_name"] == params.our_group_name)
        df_mit = df_mit[mask_serial_or_us]

    df_year = (
        df_mit[df_mit["year"] == year_to]
        .groupby("short_name")["count"]
        .sum()
        .reset_index()
        .sort_values("count", ascending=False)
        .head(params.top_n)
    )
    df_total = (
        df_mit.groupby("short_name")["count"]
        .sum()
        .reset_index()
        .sort_values("count", ascending=False)
        .head(params.top_n)
    )

    if params.our_group_name and params.our_group_name not in df_total["short_name"].values:
        our_row = df_mit[df_mit["short_name"] == params.our_group_name].groupby("short_name")["count"].sum().reset_index()
        if not our_row.empty:
            df_total = pd.concat([df_total, our_row], ignore_index=True)

    df_vri = _fetch_vri_agg(ch, year_from, year_to, params.only_applicable, mit_numbers)
    stats_year = None
    competitors = None
    if not df_vri.empty:
        df_vri["price"] = df_vri.apply(
            lambda r: _calc_verification_price(str(r["mit_number"]), str(r["mi_modification"])),
            axis=1,
        )
        df_vri["revenue"] = df_vri["price"] * df_vri["qty"]
        df_vri["is_us"] = df_vri["org_title"].str.contains(
            params.our_org_regex, case=False, regex=True, na=False
        )

        total_qty = df_vri.groupby("year")["qty"].sum()
        our_qty = df_vri[df_vri["is_us"]].groupby("year")["qty"].sum()
        our_rev = df_vri[df_vri["is_us"]].groupby("year")["revenue"].sum()
        stats_year = pd.DataFrame(
            {
                "total_qty": total_qty,
                "our_qty": our_qty,
                "our_revenue": our_rev,
            }
        ).fillna(0)
        stats_year["competitor_qty"] = stats_year["total_qty"] - stats_year["our_qty"]

        competitors = (
            df_vri[~df_vri["is_us"]]
            .groupby("org_title")
            .agg(lost_qty=("qty", "sum"), lost_money=("revenue", "sum"))
            .sort_values("lost_qty", ascending=False)
            .head(params.competitors_n)
        )

    doc = Document()
    doc.add_heading(f"Отчёт: Утверждения типов СИ и анализ поверок ({year_from}–{year_to})", level=1)
    doc.add_paragraph(f"Свой производитель/группа: {params.our_group_name}")
    doc.add_paragraph(f"Правило поиска своих (regex): {params.our_org_regex}")
    doc.add_paragraph(f"База: {ch.db}")
    doc.add_paragraph("")

    doc.add_heading("1. Утверждения типов СИ", level=1)
    if df_year.empty:
        doc.add_paragraph("Данные по утверждениям типа не найдены.")
    else:
        df_year_rep = df_year.rename(columns={"short_name": "Производитель", "count": "Кол-во типов"})
        df_year_rep.index = range(1, len(df_year_rep) + 1)
        df_year_rep.index.name = "Место"
        add_df_table(
            doc,
            df_year_rep,
            f"1.1 Лидеры по утверждению новых типов СИ ({year_to}), Топ-{params.top_n}",
            highlight_col="Производитель",
            highlight_value=params.our_group_name,
            int_cols=["Кол-во типов"],
        )

        df_total_rep = df_total.rename(columns={"short_name": "Производитель", "count": "Кол-во типов"})
        df_total_rep.index = range(1, len(df_total_rep) + 1)
        df_total_rep.index.name = "Место"
        add_df_table(
            doc,
            df_total_rep,
            f"1.2 Лидеры по утверждению типов СИ ({year_from}–{year_to}), Топ-{params.top_n}",
            highlight_col="Производитель",
            highlight_value=params.our_group_name,
            int_cols=["Кол-во типов"],
        )

        fig, ax = plt.subplots(figsize=(10, 6))
        tmp = df_total_rep.reset_index().rename(columns={"Производитель": "Производитель", "Кол-во типов": "Кол-во типов"})
        tmp = tmp.sort_values("Кол-во типов", ascending=True)
        colors = ["black" if x == params.our_group_name else "#d9d9d9" for x in tmp["Производитель"]]
        edges = ["black" if x == params.our_group_name else "gray" for x in tmp["Производитель"]]
        ax.barh(tmp["Производитель"], tmp["Кол-во типов"], color=colors, edgecolor=edges)
        ax.set_title(f"Лидеры по утверждению типов СИ ({year_from}–{year_to})")
        ax.grid(axis="x", linestyle="--", alpha=0.5)
        for y, v, name in zip(tmp["Производитель"], tmp["Кол-во типов"], tmp["Производитель"]):
            ax.text(v + 0.2, y, f"{int(v)}", va="center", fontsize=10, fontweight="bold" if name == params.our_group_name else "normal")
        fig.tight_layout()
        doc.add_paragraph("Рисунок 1 — Лидеры по утверждению типов СИ.")
        doc.add_picture(_fig_to_bytes(fig), width=Inches(6.5))
        doc.add_paragraph("")

    doc.add_heading("2. Анализ поверок", level=1)
    if stats_year is None or stats_year.empty:
        doc.add_paragraph("Данные по поверкам не найдены (по заданным фильтрам).")
    else:
        stats_rep = stats_year[["our_qty", "competitor_qty", "our_revenue"]].rename(
            columns={
                "our_qty": "Наши поверки (шт)",
                "competitor_qty": "Конкуренты (шт)",
                "our_revenue": "Наша выручка (руб)",
            }
        )
        stats_rep.index.name = "Год"
        add_df_table(
            doc,
            stats_rep,
            "2.1 Сводка по годам (мы vs рынок)",
            int_cols=["Наши поверки (шт)", "Конкуренты (шт)"],
            money_cols=["Наша выручка (руб)"],
        )

        fig, ax = plt.subplots(figsize=(10, 6))
        ax.bar(stats_year.index.astype(str), stats_year["our_qty"], label="Наши поверки", edgecolor="black")
        ax.bar(
            stats_year.index.astype(str),
            stats_year["competitor_qty"],
            bottom=stats_year["our_qty"],
            label="Упущенные (Конкуренты)",
            edgecolor="black",
            hatch="///",
            color="white",
        )
        ax.set_title("Распределение количества поверок (Мы vs Рынок)")
        ax.set_ylabel("Количество приборов (шт)")
        ax.set_xlabel("Год")
        ax.grid(axis="y", linestyle="--", alpha=0.5)
        ax.legend(frameon=False, loc="upper left")
        fig.tight_layout()
        doc.add_paragraph("Рисунок 2 — Распределение количества поверок (мы vs рынок).")
        doc.add_picture(_fig_to_bytes(fig), width=Inches(6.5))
        doc.add_paragraph("")

        fig, ax = plt.subplots(figsize=(10, 6))
        rev_mln = stats_year["our_revenue"] / 1_000_000
        bars = ax.bar(stats_year.index.astype(str), rev_mln, edgecolor="black")
        for bar in bars:
            height = bar.get_height()
            ax.text(
                bar.get_x() + bar.get_width() / 2.0,
                height + 0.2,
                f"{height:.1f} млн",
                ha="center",
                va="bottom",
                fontsize=11,
                fontweight="bold",
            )
        ax.set_title("Динамика нашей выручки от поверок")
        ax.set_ylabel("Выручка (млн руб)")
        ax.set_xlabel("Год")
        ax.grid(axis="y", linestyle="--", alpha=0.5)
        fig.tight_layout()
        doc.add_paragraph("Рисунок 3 — Динамика нашей выручки от поверок.")
        doc.add_picture(_fig_to_bytes(fig), width=Inches(6.5))
        doc.add_paragraph("")

        if competitors is not None and not competitors.empty:
            comp_rep = competitors.reset_index().rename(
                columns={"org_title": "Организация", "lost_qty": "Кол-во (шт)", "lost_money": "Оценка денег (руб)"}
            )
            comp_rep.index = range(1, len(comp_rep) + 1)
            comp_rep.index.name = "Место"
            add_df_table(
                doc,
                comp_rep,
                f"2.2 ТОП-{params.competitors_n} конкурентов (куда уходят приборы)",
                int_cols=["Кол-во (шт)"],
                money_cols=["Оценка денег (руб)"],
            )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
